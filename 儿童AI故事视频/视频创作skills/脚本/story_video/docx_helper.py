from __future__ import annotations

from pathlib import Path
import docx
from docx.shared import Inches
from .common import load_json, update_manifest

def generate_initial_docx(project_dir: Path, story_path: Path, output_docx_path: Path) -> Path:
    story = load_json(story_path)
    doc = docx.Document()
    
    doc.add_heading(story.get("title", "未命名故事"), 0)
    
    creator = story.get("creator_display", "")
    if creator:
        doc.add_paragraph(f"创作者：{creator}")
    
    doc.add_heading("角色设定", level=1)
    char = story.get("character", {})
    p = doc.add_paragraph()
    p.add_run(f"★ 名字：{char.get('name', '无')}\n")
    p.add_run(f"★ 类型：{char.get('subject_type', '无')}\n")
    p.add_run(f"★ 外貌：{char.get('appearance', '无')}\n")
    p.add_run(f"★ 性格：{char.get('personality', '无')}\n")
    p.add_run(f"★ 特殊能力：{char.get('special_ability', '无')}\n")
    
    doc.add_heading("画面风格", level=1)
    doc.add_paragraph(story.get("style", "无"))
    
    doc.add_heading("故事内容", level=1)
    for scene in story.get("scenes", []):
        idx = scene.get("index", 1)
        doc.add_paragraph(f"--- [ 第 {idx} 幕：{scene.get('title', f'场景 {idx}')} ] ---")
        doc.add_paragraph(f"【旁白-开始】\n{scene.get('narration', '')}\n【旁白-结束】")
        doc.add_paragraph(f"【画面提示词-开始】\n{scene.get('visual_prompt', '')}\n【画面提示词-结束】")
        doc.add_paragraph("【配图位置 - 待生成】")
        doc.add_paragraph("")
        
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx_path)
    update_manifest(project_dir, "docx_init", [output_docx_path])
    return output_docx_path

def update_docx_with_images(project_dir: Path, story_path: Path, output_docx_path: Path, images_dir: Path) -> Path:
    story = load_json(story_path)
    doc = docx.Document()
    
    doc.add_heading(story.get("title", "未命名故事"), 0)
    
    creator = story.get("creator_display", "")
    if creator:
        doc.add_paragraph(f"创作者：{creator}")
        
    doc.add_heading("角色设定", level=1)
    char = story.get("character", {})
    p = doc.add_paragraph()
    p.add_run(f"★ 名字：{char.get('name', '无')}\n")
    p.add_run(f"★ 类型：{char.get('subject_type', '无')}\n")
    p.add_run(f"★ 外貌：{char.get('appearance', '无')}\n")
    p.add_run(f"★ 性格：{char.get('personality', '无')}\n")
    p.add_run(f"★ 特殊能力：{char.get('special_ability', '无')}\n")
    
    doc.add_heading("画面风格", level=1)
    doc.add_paragraph(story.get("style", "无"))
    
    char_ref = images_dir / "角色参考.jpg"
    if char_ref.exists():
        doc.add_heading("角色标准参考图", level=1)
        doc.add_picture(str(char_ref), width=Inches(3.5))
        
    doc.add_heading("故事内容", level=1)
    for scene in story.get("scenes", []):
        idx = scene.get("index", 1)
        doc.add_paragraph(f"--- [ 第 {idx} 幕：{scene.get('title', f'场景 {idx}')} ] ---")
        doc.add_paragraph(f"【旁白-开始】\n{scene.get('narration', '')}\n【旁白-结束】")
        doc.add_paragraph(f"【画面提示词-开始】\n{scene.get('visual_prompt', '')}\n【画面提示词-结束】")
        
        scene_img = images_dir / f"场景_{idx:02d}.jpg"
        if scene_img.exists():
            doc.add_paragraph("【配图】")
            doc.add_picture(str(scene_img), width=Inches(3.5))
        else:
            doc.add_paragraph("【配图位置 - 缺失】")
        doc.add_paragraph("")
        
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx_path)
    update_manifest(project_dir, "docx_images", [output_docx_path])
    return output_docx_path


def parse_docx_to_story(docx_path: Path) -> dict[str, Any]:
    import re
    import docx
    doc = docx.Document(docx_path)
    story: dict[str, Any] = {
        "title": "未命名故事",
        "creator_display": "",
        "character": {
            "name": "",
            "subject_type": "object",
            "appearance": "",
            "personality": "",
            "special_ability": ""
        },
        "style": "",
        "scenes": []
    }
    
    current_section = None
    current_scene: dict[str, Any] = {}
    collecting_mode = None  # "narration" or "visual_prompt"
    collected_text = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        if p.style.name == 'Title' or p.style.name == 'Heading 0':
            story["title"] = text
            continue

        if text.startswith("创作者：") or text.startswith("创作者:"):
            parts = text.split("：", 1) if "：" in text else text.split(":", 1)
            story["creator_display"] = parts[1].strip()
            continue

        if p.style.name.startswith("Heading"):
            val = text.replace(" ", "")
            if "角色设定" in val:
                current_section = "character"
            elif "画面风格" in val:
                current_section = "style"
            elif "故事内容" in val:
                current_section = "scenes"
            continue

        if current_section == "character":
            lines = text.split("\n")
            for line in lines:
                line = line.strip()
                if "★ 名字：" in line or "★ 名字:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    story["character"]["name"] = parts[1].strip()
                elif "★ 类型：" in line or "★ 类型:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    story["character"]["subject_type"] = parts[1].strip()
                elif "★ 外貌：" in line or "★ 外貌:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    story["character"]["appearance"] = parts[1].strip()
                elif "★ 性格：" in line or "★ 性格:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    story["character"]["personality"] = parts[1].strip()
                elif "★ 特殊能力：" in line or "★ 特殊能力:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    story["character"]["special_ability"] = parts[1].strip()
            continue

        if current_section == "style":
            story["style"] = text
            current_section = None
            continue

        if current_section == "scenes" or current_section is None:
            scene_match = re.match(r"--- \[\s*第\s*(\d+)\s*幕：\s*(.*?)\s*\] ---", text)
            if scene_match:
                if current_scene:
                    story["scenes"].append(current_scene)
                current_scene = {
                    "index": int(scene_match.group(1)),
                    "title": scene_match.group(2).strip(),
                    "narration": "",
                    "visual_prompt": ""
                }
                collecting_mode = None
                collected_text = []
                continue

            if current_scene:
                if "【旁白-开始】" in text:
                    collecting_mode = "narration"
                    content = text.replace("【旁白-开始】", "").strip()
                    if "【旁白-结束】" in content:
                        content = content.replace("【旁白-结束】", "").strip()
                        current_scene["narration"] = content
                        collecting_mode = None
                    else:
                        collected_text = [content] if content else []
                    continue

                if collecting_mode == "narration":
                    if "【旁白-结束】" in text:
                        content = text.replace("【旁白-结束】", "").strip()
                        if content:
                            collected_text.append(content)
                        current_scene["narration"] = "\n".join(collected_text).strip()
                        collecting_mode = None
                        collected_text = []
                    else:
                        collected_text.append(text)
                    continue

                if "【画面提示词-开始】" in text:
                    collecting_mode = "visual_prompt"
                    content = text.replace("【画面提示词-开始】", "").strip()
                    if "【画面提示词-结束】" in content:
                        content = content.replace("【画面提示词-结束】", "").strip()
                        current_scene["visual_prompt"] = content
                        collecting_mode = None
                    else:
                        collected_text = [content] if content else []
                    continue

                if collecting_mode == "visual_prompt":
                    if "【画面提示词-结束】" in text:
                        content = text.replace("【画面提示词-结束】", "").strip()
                        if content:
                            collected_text.append(content)
                        current_scene["visual_prompt"] = "\n".join(collected_text).strip()
                        collecting_mode = None
                        collected_text = []
                    else:
                        collected_text.append(text)
                    continue

    if current_scene:
        story["scenes"].append(current_scene)

    if not story["title"] or story["title"] == "未命名故事":
        if doc.paragraphs:
            first_text = doc.paragraphs[0].text.strip()
            if first_text and not first_text.startswith("创作者："):
                story["title"] = first_text

    return story


def sync_docx_to_json(project: Path) -> None:
    docx_path = project / "故事.docx"
    if not docx_path.exists():
        return
    
    print("正在从 故事.docx 解析并同步最新内容……")
    story = parse_docx_to_story(docx_path)
    
    from .common import project_paths, save_json
    paths = project_paths(project)
    story_target = paths["text"] / "故事.json"
    save_json(story_target, story)

from __future__ import annotations
from typing import Any

from pathlib import Path
import docx
from docx.shared import Inches
from .common import load_json, update_manifest

def generate_initial_docx(project_dir: Path, story_path: Path, output_docx_path: Path) -> Path:
    story = load_json(story_path)
    doc = docx.Document()
    
    doc.add_heading(story.get("title", "未命名故事"), 0)
    
    creator = story.get("creator_display", "")
    if creator:
        doc.add_paragraph(f"创作者：{creator}")
    
    characters = story.get("characters", [])
    if not characters and "character" in story:
        characters = [story["character"]]
        
    if not characters:
        characters = [{
            "name": "角色一",
            "subject_type": "object",
            "appearance": "",
            "personality": "",
            "special_ability": ""
        }]
        
    for idx, char in enumerate(characters, 1):
        heading_text = f"角色设定 {idx}" if len(characters) > 1 else "角色设定"
        doc.add_heading(heading_text, level=1)
        p = doc.add_paragraph()
        p.add_run(f"★ 名字：{char.get('name', '无')}\n")
        p.add_run(f"★ 类型：{char.get('subject_type', '无')}\n")
        p.add_run(f"★ 外貌：{char.get('appearance', '无')}\n")
        p.add_run(f"★ 性格：{char.get('personality', '无')}\n")
        p.add_run(f"★ 特殊能力：{char.get('special_ability', '无')}\n")
    
    doc.add_heading("配音设定", level=1)
    p = doc.add_paragraph()
    p.add_run(f"★ 配音音色：{story.get('voice', '冰糖')} (仅在无个性化录音时生效，可选官方音色: 冰糖、苏打、茉莉、白桦)\n")
    
    doc.add_heading("画面风格", level=1)
    doc.add_paragraph(story.get("style", "无"))
    
    doc.add_heading("故事内容", level=1)
    for scene in story.get("scenes", []):
        idx = scene.get("index", 1)
        doc.add_paragraph(f"--- [ 第 {idx} 幕：{scene.get('title', f'场景 {idx}')} ] ---")
        doc.add_paragraph(f"【旁白-开始】\n{scene.get('narration', '')}\n【旁白-结束】")
        doc.add_paragraph(f"【画面提示词-开始】\n{scene.get('visual_prompt', '')}\n【画面提示词-结束】")
        doc.add_paragraph("【配图位置 - 待生成】")
        doc.add_paragraph("")
        
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx_path)
    update_manifest(project_dir, "docx_init", [output_docx_path])
    return output_docx_path

def update_docx_with_images(project_dir: Path, story_path: Path, output_docx_path: Path, images_dir: Path) -> Path:
    story = load_json(story_path)
    doc = docx.Document()
    
    doc.add_heading(story.get("title", "未命名故事"), 0)
    
    creator = story.get("creator_display", "")
    if creator:
        doc.add_paragraph(f"创作者：{creator}")
        
    characters = story.get("characters", [])
    if not characters and "character" in story:
        characters = [story["character"]]
        
    for idx, char in enumerate(characters, 1):
        heading_text = f"角色设定 {idx}" if len(characters) > 1 else "角色设定"
        doc.add_heading(heading_text, level=1)
        p = doc.add_paragraph()
        p.add_run(f"★ 名字：{char.get('name', '无')}\n")
        p.add_run(f"★ 类型：{char.get('subject_type', '无')}\n")
        p.add_run(f"★ 外貌：{char.get('appearance', '无')}\n")
        p.add_run(f"★ 性格：{char.get('personality', '无')}\n")
        p.add_run(f"★ 特殊能力：{char.get('special_ability', '无')}\n")
        
        char_ref_new = images_dir / f"角色参考_{idx}.jpg"
        char_ref_old = images_dir / "角色参考.jpg"
        char_ref = char_ref_new if char_ref_new.exists() else char_ref_old
        if char_ref.exists() and (idx == 1 or char_ref == char_ref_new):
            doc.add_paragraph("【角色标准参考图】")
            doc.add_picture(str(char_ref), width=Inches(3.5))
        
    doc.add_heading("配音设定", level=1)
    p = doc.add_paragraph()
    p.add_run(f"★ 配音音色：{story.get('voice', '冰糖')} (仅在无个性化录音时生效，可选官方音色: 冰糖、苏打、茉莉、白桦)\n")
    
    doc.add_heading("画面风格", level=1)
    doc.add_paragraph(story.get("style", "无"))
        
    doc.add_heading("故事内容", level=1)
    for scene in story.get("scenes", []):
        idx = scene.get("index", 1)
        doc.add_paragraph(f"--- [ 第 {idx} 幕：{scene.get('title', f'场景 {idx}')} ] ---")
        doc.add_paragraph(f"【旁白-开始】\n{scene.get('narration', '')}\n【旁白-结束】")
        doc.add_paragraph(f"【画面提示词-开始】\n{scene.get('visual_prompt', '')}\n【画面提示词-结束】")
        
        scene_img = images_dir / f"场景_{idx:02d}.jpg"
        if scene_img.exists():
            doc.add_paragraph("【配图】")
            doc.add_picture(str(scene_img), width=Inches(3.5))
        else:
            doc.add_paragraph("【配图位置 - 缺失】")
        doc.add_paragraph("")
        
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx_path)
    update_manifest(project_dir, "docx_images", [output_docx_path])
    return output_docx_path


def parse_docx_to_story(docx_path: Path) -> dict[str, Any]:
    import re
    import docx
    doc = docx.Document(docx_path)
    story: dict[str, Any] = {
        "title": "未命名故事",
        "creator_display": "",
        "character": {
            "name": "",
            "subject_type": "object",
            "appearance": "",
            "personality": "",
            "special_ability": ""
        },
        "characters": [],
        "style": "",
        "scenes": []
    }
    
    current_section = None
    current_character: dict[str, Any] = {}
    characters: list[dict[str, Any]] = []
    current_scene: dict[str, Any] = {}
    collecting_mode = None  # "narration" or "visual_prompt"
    collected_text = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
            
        if p.style.name == 'Title' or p.style.name == 'Heading 0':
            story["title"] = text
            continue

        creator_match = re.match(r"^(?:创作者|作者|制作人|作者姓名)\s*[：:]\s*(.*)$", text)
        if creator_match:
            story["creator_display"] = creator_match.group(1).strip()
            continue

        if p.style.name.startswith("Heading"):
            val = text.replace(" ", "")
            if re.match(r"角色设定\s*(\d+)?", val):
                current_section = "character"
                if current_character:
                    characters.append(current_character)
                current_character = {
                    "name": "",
                    "subject_type": "object",
                    "appearance": "",
                    "personality": "",
                    "special_ability": ""
                }
            elif "配音设定" in val:
                current_section = "voice"
                if current_character:
                    characters.append(current_character)
                    current_character = {}
            elif "画面风格" in val:
                current_section = "style"
                if current_character:
                    characters.append(current_character)
                    current_character = {}
            elif "故事内容" in val:
                current_section = "scenes"
                if current_character:
                    characters.append(current_character)
                    current_character = {}
            continue

        if current_section == "character" and current_character is not None:
            lines = text.split("\n")
            for line in lines:
                line = line.strip()
                if "★ 名字：" in line or "★ 名字:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    current_character["name"] = parts[1].strip()
                elif "★ 类型：" in line or "★ 类型:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    current_character["subject_type"] = parts[1].strip()
                elif "★ 外貌：" in line or "★ 外貌:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    current_character["appearance"] = parts[1].strip()
                elif "★ 性格：" in line or "★ 性格:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    current_character["personality"] = parts[1].strip()
                elif "★ 特殊能力：" in line or "★ 特殊能力:" in line:
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    current_character["special_ability"] = parts[1].strip()
            continue

        if current_section == "voice":
            lines = text.split("\n")
            for line in lines:
                line = line.strip()
                if any(k in line for k in ["★ 配音音色：", "★ 配音音色:", "★ 声音：", "★ 声音:"]):
                    parts = line.split("：", 1) if "：" in line else line.split(":", 1)
                    val = parts[1].strip()
                    val = re.split(r'[\(\s]', val)[0].strip()
                    story["voice"] = val
            continue

        if current_section == "style":
            story["style"] = text
            current_section = None
            continue

        if current_section == "scenes" or current_section is None:
            scene_match = re.match(r"--- \[\s*第\s*(\d+)\s*幕：\s*(.*?)\s*\] ---", text)
            if scene_match:
                if current_scene:
                    story["scenes"].append(current_scene)
                current_scene = {
                    "index": int(scene_match.group(1)),
                    "title": scene_match.group(2).strip(),
                    "narration": "",
                    "visual_prompt": ""
                }
                collecting_mode = None
                collected_text = []
                continue

            if current_scene:
                if "【旁白-开始】" in text:
                    collecting_mode = "narration"
                    content = text.replace("【旁白-开始】", "").strip()
                    if "【旁白-结束】" in content:
                        content = content.replace("【旁白-结束】", "").strip()
                        current_scene["narration"] = content
                        collecting_mode = None
                    else:
                        collected_text = [content] if content else []
                    continue

                if collecting_mode == "narration":
                    if "【旁白-结束】" in text:
                        content = text.replace("【旁白-结束】", "").strip()
                        if content:
                            collected_text.append(content)
                        current_scene["narration"] = "\n".join(collected_text).strip()
                        collecting_mode = None
                        collected_text = []
                    else:
                        collected_text.append(text)
                    continue

                if "【画面提示词-开始】" in text:
                    collecting_mode = "visual_prompt"
                    content = text.replace("【画面提示词-开始】", "").strip()
                    if "【画面提示词-结束】" in content:
                        content = content.replace("【画面提示词-结束】", "").strip()
                        current_scene["visual_prompt"] = content
                        collecting_mode = None
                    else:
                        collected_text = [content] if content else []
                    continue

                if collecting_mode == "visual_prompt":
                    if "【画面提示词-结束】" in text:
                        content = text.replace("【画面提示词-结束】", "").strip()
                        if content:
                            collected_text.append(content)
                        current_scene["visual_prompt"] = "\n".join(collected_text).strip()
                        collecting_mode = None
                        collected_text = []
                    else:
                        collected_text.append(text)
                    continue

    if current_character:
        characters.append(current_character)
    if current_scene:
        story["scenes"].append(current_scene)

    story["characters"] = [c for c in characters if c.get("name")]
    if story["characters"]:
        story["character"] = story["characters"][0]

    if not story["title"] or story["title"] == "未命名故事":
        if doc.paragraphs:
            first_text = doc.paragraphs[0].text.strip()
            if first_text and not first_text.startswith("创作者："):
                story["title"] = first_text

    return story


def sync_docx_to_json(project: Path) -> None:
    docx_path = project / "故事.docx"
    if not docx_path.exists():
        return
    
    print("正在从 故事.docx 解析并同步最新内容……")
    story = parse_docx_to_story(docx_path)
    
    from .common import project_paths, save_json
    paths = project_paths(project)
    story_target = paths["text"] / "故事.json"
    save_json(story_target, story)
